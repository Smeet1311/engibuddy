from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

doc.add_heading('Push-Back Mechanism — Test Results', level=0)

doc.add_paragraph(
    'Goal: make sure a student can never skip ahead of an unfinished step, '
    'in both Guidance Mode and Review Mode. Each of the 4 scenarios below was '
    'tried, including ways to break it on purpose.'
)

doc.add_heading('Overall Result', level=1)
p = doc.add_paragraph()
run = p.add_run('ALL 4 SCENARIOS: PASS')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1a, 0x7a, 0x1a)

doc.add_paragraph(
    '39 automated tests pass, plus the same 4 scenarios were also run for real '
    'against the live app and a real AI model (not just simulated) — both checks agree.'
)

scenarios = [
    (
        '1. Software development project',
        "Tried saying the whole app was already built and deployed, skipping "
        "straight to the end, while step 1 was still empty.",
        'PASS',
        'App correctly refused the shortcut and kept the student on step 1 until '
        'it was actually done. Also checked the opposite: once steps 1-3 are '
        'genuinely finished, moving to step 4 is allowed right away — no false '
        'blocking either.',
    ),
    (
        '2. Bridge building + document upload',
        'Uploaded a document that looked like it finished a later step '
        '(design/build), while the very first step was still untouched.',
        'PASS',
        "The document's content was accepted and credited, but the app still "
        "would not let the student skip the unfinished first step. Also "
        "confirmed: removing an uploaded document correctly undoes the credit "
        "it gave.",
    ),
    (
        '3. Starting a project with no idea',
        "Sent a vague message (\"I don't have an idea yet\") and also tried to "
        "trick the app with confusing/garbled input.",
        'PASS',
        'App stayed calmly on step 1, asked a helpful follow-up question, and '
        'never crashed or got confused, even with garbled input.',
    ),
    (
        '4. Project already 2 weeks in, with shared documents (Review Mode check-in)',
        'Built a real project with 2 weeks of real progress, then opened the '
        'Review tab and said "we are completely done, can you confirm?" — '
        'trying to get Review Mode to just believe it.',
        'PASS',
        'Review Mode correctly checked the real progress and held the student '
        'at the true current step instead of believing the claim. (This used '
        'to fail before our fix — Review Mode had no check at all. Now fixed '
        'and confirmed working for real.)',
    ),
]

doc.add_heading('Scenario Results', level=1)
for name, attempt, result, outcome in scenarios:
    doc.add_heading(name, level=2)
    t = doc.add_table(rows=3, cols=2)
    t.style = 'Light Grid Accent 1'
    t.rows[0].cells[0].text = 'What we tried to break'
    t.rows[0].cells[1].text = attempt
    t.rows[1].cells[0].text = 'Result'
    cell = t.rows[1].cells[1]
    cell.text = ''
    r = cell.paragraphs[0].add_run(result)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1a, 0x7a, 0x1a)
    t.rows[2].cells[0].text = 'What actually happened'
    t.rows[2].cells[1].text = outcome
    doc.add_paragraph()

doc.add_heading('Issues Found Along the Way (all fixed)', level=1)
doc.add_paragraph(
    'While testing, we found a few real problems in the app and fixed all of '
    'them before finishing:'
)
issues = [
    'Review Mode had no safety check at all — fixed so it now checks real '
    'progress just like Guidance Mode does.',
    'If a student finished every step and the AI got confused about which '
    'step came next, the app could crash — fixed so it never goes out of bounds.',
    'The Review tab could briefly show the wrong (empty) checklist after '
    'sending a message — fixed so it always shows the real, correct progress.',
]
for i in issues:
    doc.add_paragraph(i, style='List Bullet')

doc.add_heading('How to re-check this', level=1)
doc.add_paragraph('Run: cd backend && python -m pytest tests/test_pushback_scenarios.py -v')
doc.add_paragraph('All 14 automated tests for these scenarios should show PASSED.')

doc.save('D:/engibuddy-merge/PUSHBACK_TEST_RESULTS.docx')
print('saved')
